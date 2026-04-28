"""DOCX 解析器测试。

构造一个 fixture .docx，验证：
- 标题/段落/表格按 XML 顺序输出
- 标题级别识别
- 表格被正确识别为 TABLE block
- paragraph_index 全局递增
- anchor_text 非空且来自原文开头
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.parse.docx_parser import DocxBlockType, parse_docx


@pytest.fixture
def fixture_docx(tmp_path: Path) -> Path:
    """构造一份小的 .docx 用作测试。"""
    p = tmp_path / "fixture.docx"
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("这是第一段，简介内容。")
    doc.add_paragraph("第二段，继续介绍。")
    doc.add_heading("1.1 子章节", level=2)
    doc.add_paragraph("子章节段落。")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "值"
    table.rows[1].cells[0].text = "员工数"
    table.rows[1].cells[1].text = "120"
    table.rows[2].cells[0].text = "里程"
    table.rows[2].cells[1].text = "85"
    doc.add_heading("第二章 操作", level=1)
    doc.add_paragraph("操作流程。")
    doc.save(p)
    return p


def test_parse_yields_blocks(fixture_docx: Path) -> None:
    blocks = parse_docx(fixture_docx)
    assert len(blocks) >= 6


def test_heading_levels(fixture_docx: Path) -> None:
    blocks = parse_docx(fixture_docx)
    headings = [b for b in blocks if b.block_type == DocxBlockType.HEADING]
    assert len(headings) == 3
    assert headings[0].heading_level == 1
    assert headings[0].text.startswith("第一章")
    assert headings[1].heading_level == 2
    assert headings[2].heading_level == 1


def test_paragraph_index_monotonic(fixture_docx: Path) -> None:
    blocks = parse_docx(fixture_docx)
    indices = [b.paragraph_index for b in blocks]
    assert indices == sorted(indices)
    assert indices == sorted(set(indices))


def test_table_extracted(fixture_docx: Path) -> None:
    blocks = parse_docx(fixture_docx)
    tables = [b for b in blocks if b.block_type == DocxBlockType.TABLE]
    assert len(tables) == 1
    rows = tables[0].table_rows
    assert len(rows) == 3
    assert rows[0] == ["字段", "值"]
    assert rows[1] == ["员工数", "120"]


def test_anchor_text_present(fixture_docx: Path) -> None:
    blocks = parse_docx(fixture_docx)
    for b in blocks:
        if b.block_type in (DocxBlockType.HEADING, DocxBlockType.PARAGRAPH):
            assert b.anchor_text, f"{b!r} 缺少 anchor_text"


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        parse_docx(tmp_path / "nope.docx")


@pytest.fixture
def pseudo_heading_docx(tmp_path: Path) -> Path:
    """工业文档常见：用普通段落写编号充当标题。"""
    p = tmp_path / "pseudo.docx"
    doc = Document()
    doc.add_paragraph("第一章 概况")
    doc.add_paragraph("本文档介绍。")
    doc.add_paragraph("一、岗位条件")
    doc.add_paragraph("具体要求。")
    doc.add_paragraph("1.1 培训")
    doc.add_paragraph("培训内容。")
    doc.add_paragraph("（一）应急处置")
    doc.add_paragraph("步骤如下。")
    doc.add_paragraph("附录 A 操作票模板")
    doc.save(p)
    return p


def test_pseudo_heading_detected(pseudo_heading_docx: Path) -> None:
    blocks = parse_docx(pseudo_heading_docx)
    headings = [b for b in blocks if b.block_type == DocxBlockType.HEADING]
    titles = [h.text for h in headings]
    assert any("第一章" in t for t in titles)
    assert any("岗位条件" in t for t in titles)
    assert any("1.1" in t for t in titles)
    assert any("应急处置" in t for t in titles)
    assert any("附录" in t for t in titles)


def test_pseudo_heading_levels(pseudo_heading_docx: Path) -> None:
    """1.1 应该是 level 2；第一章 / 一、 应该是 level 1。"""
    blocks = parse_docx(pseudo_heading_docx)
    by_text = {b.text: b for b in blocks if b.block_type == DocxBlockType.HEADING}
    assert by_text["第一章 概况"].heading_level == 1
    assert by_text["一、岗位条件"].heading_level == 1
    assert by_text["1.1 培训"].heading_level == 2
    assert by_text["（一）应急处置"].heading_level == 3


def test_long_paragraph_not_heading(tmp_path: Path) -> None:
    """虽然以"1."开头但太长（>40字），不应识别为标题。"""
    p = tmp_path / "long.docx"
    doc = Document()
    doc.add_paragraph("1. 这是一段很长的正文内容包含了很多说明文字以及操作步骤详细描述本应该是普通段落不是标题")
    doc.save(p)
    blocks = parse_docx(p)
    paragraphs = [b for b in blocks if b.block_type == DocxBlockType.PARAGRAPH]
    assert len(paragraphs) == 1


@pytest.mark.integration
def test_parse_real_sample(sample_docx_path: Path, tmp_path: Path) -> None:
    """真实样本：一区一案，要能解出标题和表格。"""
    blocks = parse_docx(sample_docx_path, image_out_dir=tmp_path / "img")
    assert len(blocks) > 50
    # 一区一案样本里至少有几个标题 + 多个表格
    assert any(b.block_type == DocxBlockType.TABLE for b in blocks)
    # 应该提取到图片
    images = [b for b in blocks if b.block_type == DocxBlockType.IMAGE]
    assert len(images) > 0
