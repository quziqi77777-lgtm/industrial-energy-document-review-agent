"""切块器测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.chunk import ChunkType, chunk_docx_blocks
from src.parse.docx_parser import parse_docx


@pytest.fixture
def hierarchical_docx(tmp_path: Path) -> Path:
    p = tmp_path / "hier.docx"
    doc = Document()
    doc.add_heading("1 概述", level=1)
    doc.add_paragraph("总体介绍。")
    doc.add_heading("1.1 目的", level=2)
    doc.add_paragraph("为了规范作业流程。")
    doc.add_paragraph("具体内容如下。")
    doc.add_heading("2 操作", level=1)
    doc.add_paragraph("按步骤执行。")
    doc.save(p)
    return p


@pytest.fixture
def with_table_docx(tmp_path: Path) -> Path:
    p = tmp_path / "table.docx"
    doc = Document()
    doc.add_heading("5 应急处置", level=1)
    doc.add_paragraph("应急步骤如下。")
    table = doc.add_table(rows=15, cols=3)
    for i in range(15):
        for j in range(3):
            table.rows[i].cells[j].text = f"r{i}c{j}"
    doc.save(p)
    return p


def test_chunk_id_format(hierarchical_docx: Path) -> None:
    blocks = parse_docx(hierarchical_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST")
    assert all(c.chunk_id.startswith("TEST__") for c in chunks)
    parts = chunks[0].chunk_id.split("__")
    assert len(parts) == 4
    assert parts[3].isdigit()


def test_section_path_inferred(hierarchical_docx: Path) -> None:
    blocks = parse_docx(hierarchical_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST")
    sections = {c.section_path for c in chunks if c.chunk_type != ChunkType.HEADING}
    assert "1.1" in sections
    assert "2" in sections


def test_chunk_id_unique(hierarchical_docx: Path) -> None:
    blocks = parse_docx(hierarchical_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST")
    ids = [c.chunk_id for c in chunks]
    assert len(ids) == len(set(ids))


def test_large_table_split(with_table_docx: Path) -> None:
    blocks = parse_docx(with_table_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST", table_inline_rows=10)
    summaries = [c for c in chunks if c.chunk_type == ChunkType.TABLE_SUMMARY]
    fulls = [c for c in chunks if c.chunk_type == ChunkType.TABLE_FULL]
    assert len(summaries) == 1
    assert len(fulls) == 1
    assert "共 15 行" in summaries[0].content
    assert summaries[0].cross_refs == [fulls[0].chunk_id]
    assert fulls[0].parent_id == summaries[0].chunk_id


def test_small_table_inlined(tmp_path: Path) -> None:
    p = tmp_path / "small.docx"
    doc = Document()
    doc.add_heading("章节", level=1)
    table = doc.add_table(rows=3, cols=2)
    for i in range(3):
        for j in range(2):
            table.rows[i].cells[j].text = f"x{i}{j}"
    doc.save(p)
    blocks = parse_docx(p)
    chunks = chunk_docx_blocks(blocks, doc_id="T", table_inline_rows=10)
    summaries = [c for c in chunks if c.chunk_type == ChunkType.TABLE_SUMMARY]
    fulls = [c for c in chunks if c.chunk_type == ChunkType.TABLE_FULL]
    assert len(summaries) == 1
    assert len(fulls) == 0


def test_paragraph_index_preserved(hierarchical_docx: Path) -> None:
    blocks = parse_docx(hierarchical_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST")
    indices = [c.paragraph_index for c in chunks]
    assert indices == sorted(indices)


def test_anchor_text_nonempty(hierarchical_docx: Path) -> None:
    blocks = parse_docx(hierarchical_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST")
    for c in chunks:
        assert c.anchor_text, f"{c.chunk_id} 缺少 anchor"


def test_to_row_serializable(hierarchical_docx: Path) -> None:
    """Chunk.to_row() 输出能直接喂给 sqlite executemany。"""
    import json

    blocks = parse_docx(hierarchical_docx)
    chunks = chunk_docx_blocks(blocks, doc_id="TEST")
    rows = [c.to_row() for c in chunks]
    for r in rows:
        json.loads(r["dimensions"])
        json.loads(r["cross_refs"])
        json.loads(r["extra"])
