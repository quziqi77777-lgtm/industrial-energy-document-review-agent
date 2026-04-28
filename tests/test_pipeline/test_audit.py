"""审核流水线测试。

只用 mock LLM；真实样本走 .doc → .docx 转换 → 切块 → E1 → JSON。
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from src.config import load_config
from src.llm.mock_provider import MockProvider
from src.pipeline.audit import (
    AuditPipeline,
    derive_doc_id,
    is_relevant_for_e1,
)
from src.chunk.models import Chunk, ChunkType
from src.store import Repository


def make_test_config(tmp_path: Path) -> dict:
    """构造一份指向 tmp_path 的配置。"""
    return {
        "llm": {
            "provider": "mock",
            "text_model": "test-model",
            "vision_model": "test-vision",
            "default_temperature": 0.0,
            "explorer_a_temperature": 0.2,
            "explorer_b_temperature": 0.0,
            "mock": {"text_response": "{}"},
        },
        "paths": {
            "data_dir": str(tmp_path),
            "db_path": ":memory:",
            "docs_dir": str(tmp_path / "docs"),
            "images_dir": str(tmp_path / "images"),
            "results_dir": str(tmp_path / "results"),
            "fixtures_dir": str(tmp_path / "fixtures"),
        },
        "parse": {
            "doc_to_docx_timeout": 60.0,
            "pdf_scan_text_threshold": 50,
            "ocr_enabled": False,
        },
        "chunk": {
            "max_tokens": 800,
            "table_inline_rows": 10,
            "anchor_text_length": 30,
        },
        "retrieve": {"fts_top_k": 3, "use_jieba": True},
        "audit": {"per_doc_timeout": 240, "batch_concurrency": 5},
    }


class TestHelpers:
    def test_derive_doc_id_basic(self) -> None:
        assert derive_doc_id(Path("/x/foo bar.doc")) == "foo_bar"

    def test_derive_doc_id_chinese(self) -> None:
        out = derive_doc_id(Path("AS作业区作业指导书(1).doc"))
        assert "_" in out
        assert len(out) <= 80

    def test_is_relevant_for_e1(self) -> None:
        c = Chunk(
            chunk_id="x", doc_id="d", chunk_type=ChunkType.TEXT,
            section_path="3", title="t",
            content="员工 34 人，巡护 69 名。",
            paragraph_index=0, anchor_text="x",
        )
        assert is_relevant_for_e1(c) is True

    def test_is_irrelevant(self) -> None:
        c = Chunk(
            chunk_id="x", doc_id="d", chunk_type=ChunkType.TEXT,
            section_path="3", title="t",
            content="天气很好",
            paragraph_index=0, anchor_text="x",
        )
        assert is_relevant_for_e1(c) is False


class TestAuditPipeline:
    @pytest.fixture
    def docx_with_staffing(self, tmp_path: Path) -> Path:
        from docx import Document

        p = tmp_path / "fake_guide.docx"
        doc = Document()
        doc.add_heading("第一章 概况", level=1)
        doc.add_paragraph("ZZ 分公司管辖天然气管道 200 公里。")
        doc.add_paragraph("公司共有员工 34 人。")
        doc.add_heading("第三章 岗位", level=1)
        doc.add_paragraph("HSE 管理工程师 1 名。")
        doc.add_paragraph("配置巡护人员 60 名覆盖全线。")
        doc.save(p)
        return p

    def test_run_with_mock_llm(self, docx_with_staffing: Path, tmp_path: Path) -> None:
        cfg = make_test_config(tmp_path)
        provider = MockProvider(
            text_response=json.dumps(
                {
                    "total_employees": 34,
                    "safety_engineers_actual": 1,
                    "pipeline_length_km": 200,
                    "pipeline_kinds": ["natural_gas"],
                    "patrol_workers_actual": 60,
                    "pipeline_count": 1,
                    "pipeline_names": ["天然气干线"],
                }
            )
        )
        pipe = AuditPipeline(config=cfg, provider=provider)
        try:
            result = pipe.run(docx_with_staffing)
        finally:
            pipe.close()

        assert result.doc_id
        assert "E1_staffing" in result.dimensions
        e1 = result.dimensions["E1_staffing"]
        assert e1["extra"]["staffing_analysis"]["total_employees"] == 34
        assert e1["extra"]["staffing_analysis"]["patrol_workers"] == 60
        # 200 km / 60 workers ≈ 3.33 km/p ∈ [3,10]
        assert e1["extra"]["staffing_analysis"]["km_per_worker"] is not None

    def test_no_relevant_chunks_returns_uncertain(self, tmp_path: Path) -> None:
        from docx import Document

        p = tmp_path / "irrelevant.docx"
        doc = Document()
        doc.add_heading("天气", level=1)
        doc.add_paragraph("今天天气很好。")
        doc.save(p)

        cfg = make_test_config(tmp_path)
        provider = MockProvider(text_response="{}")
        pipe = AuditPipeline(config=cfg, provider=provider)
        try:
            result = pipe.run(p)
        finally:
            pipe.close()
        assert result.dimensions["E1_staffing"]["verdict"] == "uncertain"
        assert result.need_human_review is True


class TestPersistence:
    def test_audit_writes_label_with_pipeline_audit(
        self, tmp_path: Path
    ) -> None:
        from docx import Document

        p = tmp_path / "g.docx"
        doc = Document()
        doc.add_paragraph("员工 34 人。")
        doc.add_paragraph("巡护 60 名。")
        doc.add_paragraph("天然气管道 200 公里。")
        doc.save(p)

        cfg = make_test_config(tmp_path)
        repo = Repository(":memory:")
        provider = MockProvider(text_response="{}")
        pipe = AuditPipeline(config=cfg, provider=provider, repo=repo)
        try:
            result = pipe.run(p)
            labels = repo.get_labels(result.doc_id, pipeline="audit")
        finally:
            pipe.close()
        # 现在 pipeline 跑 C1 + C4 + E1 + T1 + T2 + T3 共 6 个维度
        assert len(labels) == 6
        dims = {l["dimension"] for l in labels}
        assert dims == {
            "C1_structure", "C4_reference", "E1_staffing",
            "T1_template", "T2_format", "T3_latency",
        }
        assert all(l["pipeline"] == "audit" for l in labels)
        assert all(l["human_signoff"] == 0 for l in labels)


@pytest.mark.integration
def test_real_doc_end_to_end(sample_doc_path: Path, tmp_path: Path) -> None:
    """真实 .doc 跑全链路。"""
    cfg = make_test_config(tmp_path)
    provider = MockProvider(text_response="{}")
    pipe = AuditPipeline(config=cfg, provider=provider)
    try:
        result = pipe.run(sample_doc_path)
    finally:
        pipe.close()
    e1 = result.dimensions["E1_staffing"]
    analysis = e1["extra"]["staffing_analysis"]
    # 样本数据：员工 34 / 管道 274km / 巡线 69
    # MockProvider 返回 {} → 全靠 regex 兜底
    assert analysis["total_employees"] == 34
    assert analysis["patrol_workers"] == 69
    # 至少包含一种管道类型
    assert analysis["pipeline_kinds"]
